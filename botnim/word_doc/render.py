"""Markdown → docx renderer for the unified bot's Word-doc tool."""
from __future__ import annotations

import io
import re
from typing import List

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.text.parfmt import ParagraphFormat

from .markdown_walker import parse_markdown
from .models import WordDocRequest, WordDocSection


# python-docx 1.x ships no ParagraphFormat.bidi descriptor. We patch one in
# at import time so the canonical RTL toggle (`paragraph_format.bidi = True`)
# both round-trips on the Python side and emits a `<w:bidi w:val="1"/>` into
# `<w:pPr>` — which is what Word actually consumes for paragraph-level RTL.
def _bidi_get(self):
    pPr = self._element.pPr
    if pPr is None:
        return None
    bidi_el = pPr.find(qn("w:bidi"))
    if bidi_el is None:
        return None
    val = bidi_el.get(qn("w:val"))
    if val in (None, "1", "true", "True"):
        return True
    return False


def _bidi_set(self, value):
    pPr = self._element.get_or_add_pPr()
    for existing in pPr.findall(qn("w:bidi")):
        pPr.remove(existing)
    if value:
        bidi_el = OxmlElement("w:bidi")
        bidi_el.set(qn("w:val"), "1")
        pPr.append(bidi_el)


if not isinstance(getattr(ParagraphFormat, "bidi", None), property):
    ParagraphFormat.bidi = property(_bidi_get, _bidi_set)

_FILENAME_BAD_CHARS = re.compile(r"[\x00-\x1f/\\]")
_DEFAULT_FILENAME = "botnim-document.docx"


def sanitize_filename(name: str) -> str:
    name = _FILENAME_BAD_CHARS.sub("", (name or "")).strip()
    if not name:
        return _DEFAULT_FILENAME
    if not name.lower().endswith(".docx"):
        name = f"{name}.docx"
    if len(name) > 100:
        # Preserve .docx suffix; cap at exactly 100 chars (95 + ".docx").
        name = name[:95] + ".docx"
    return name


def _force_rtl(paragraph):
    paragraph.paragraph_format.bidi = True


def _force_rtl_run(run):
    run.font.rtl = True
    rpr = run._element.get_or_add_rPr()
    bidi = OxmlElement("w:rtl")
    bidi.set(qn("w:val"), "1")
    # Remove any prior <w:rtl> to keep it idempotent
    for existing in rpr.findall(qn("w:rtl")):
        rpr.remove(existing)
    rpr.append(bidi)


def _setup_styles(doc):
    styles = doc.styles
    for name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"):
        try:
            style = styles[name]
        except KeyError:
            continue
        style.paragraph_format.bidi = True
        if hasattr(style, "font"):
            style.font.rtl = True
            try:
                style.font.name = "David"
            except Exception:
                pass


def _add_hyperlink(paragraph, url: str, text: str):
    """python-docx has no high-level hyperlink API, so we insert the OXML by hand."""
    part = paragraph.part
    rid = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    rpr.append(rtl)
    new_run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _render_runs(paragraph, runs: List[dict]):
    for r in runs:
        if r.get("url"):
            _add_hyperlink(paragraph, r["url"], r["text"])
        else:
            run = paragraph.add_run(r["text"])
            run.bold = bool(r["bold"])
            run.italic = bool(r["italic"])
            _force_rtl_run(run)


def _render_section(doc, section: WordDocSection):
    heading_para = doc.add_heading(section.heading, level=section.level)
    _force_rtl(heading_para)
    for run in heading_para.runs:
        _force_rtl_run(run)

    blocks = parse_markdown(section.body_md)
    for block in blocks:
        if block["type"] == "heading":
            sub = doc.add_heading(block["text"], level=min(block["level"] + section.level, 9))
            _force_rtl(sub)
            for r in sub.runs:
                _force_rtl_run(r)
        elif block["type"] == "list_item":
            style = "List Number" if block["ordered"] else "List Bullet"
            p = doc.add_paragraph(style=style)
            _force_rtl(p)
            _render_runs(p, block["runs"])
        else:  # paragraph
            p = doc.add_paragraph()
            _force_rtl(p)
            _render_runs(p, block["runs"])


def render_word_doc(req: WordDocRequest) -> bytes:
    doc = Document()
    _setup_styles(doc)

    # Title page
    title_para = doc.add_heading(req.title, level=0)
    _force_rtl(title_para)
    for r in title_para.runs:
        _force_rtl_run(r)

    for section in req.sections:
        _render_section(doc, section)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
