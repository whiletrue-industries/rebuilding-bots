"""Tests for the markdown → .docx renderer."""
from __future__ import annotations

import io
import pytest
from docx import Document

from botnim.word_doc.render import render_word_doc, sanitize_filename
from botnim.word_doc.models import WordDocRequest, WordDocSection


def _docx_bytes_to_doc(b: bytes):
    return Document(io.BytesIO(b))


def test_renders_title_and_one_section():
    req = WordDocRequest(
        title="כותרת",
        sections=[WordDocSection(heading="רקע", level=1, body_md="פסקה ראשונה")],
    )
    out = render_word_doc(req)
    assert isinstance(out, bytes)
    doc = _docx_bytes_to_doc(out)
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert "כותרת" in headings or any("כותרת" in p.text for p in doc.paragraphs[:3])
    assert any("רקע" == p.text for p in doc.paragraphs)
    assert any("פסקה ראשונה" in p.text for p in doc.paragraphs)


def test_every_paragraph_is_rtl():
    req = WordDocRequest(
        title="t",
        sections=[WordDocSection(heading="h", level=1, body_md="a\n\nb\n\nc")],
    )
    out = render_word_doc(req)
    doc = _docx_bytes_to_doc(out)
    for p in doc.paragraphs:
        assert p.paragraph_format.bidi is True, f"paragraph {p.text!r} not RTL"


def test_bullet_list_renders():
    req = WordDocRequest(
        title="t",
        sections=[WordDocSection(heading="h", level=1, body_md="- one\n- two\n- three")],
    )
    out = render_word_doc(req)
    doc = _docx_bytes_to_doc(out)
    list_paras = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
    assert len(list_paras) == 3


def test_numbered_list_renders():
    req = WordDocRequest(
        title="t",
        sections=[WordDocSection(heading="h", level=1, body_md="1. one\n2. two")],
    )
    out = render_word_doc(req)
    doc = _docx_bytes_to_doc(out)
    list_paras = [p for p in doc.paragraphs if p.style.name == "List Number"]
    assert len(list_paras) == 2


def test_bold_run_emitted():
    req = WordDocRequest(
        title="t",
        sections=[WordDocSection(heading="h", level=1, body_md="hello **world** end")],
    )
    out = render_word_doc(req)
    doc = _docx_bytes_to_doc(out)
    found = False
    for p in doc.paragraphs:
        for r in p.runs:
            if r.text == "world" and r.bold:
                found = True
    assert found


def test_italic_run_emitted():
    req = WordDocRequest(
        title="t",
        sections=[WordDocSection(heading="h", level=1, body_md="hello *world* end")],
    )
    out = render_word_doc(req)
    doc = _docx_bytes_to_doc(out)
    found = any(r.italic for p in doc.paragraphs for r in p.runs if r.text == "world")
    assert found


def test_hyperlink_emitted():
    req = WordDocRequest(
        title="t",
        sections=[WordDocSection(heading="h", level=1, body_md="see [docs](https://x.example) here")],
    )
    out = render_word_doc(req)
    # docx is a deflate-compressed zip; the hyperlink Target lives in
    # word/_rels/document.xml.rels. Extract and assert there.
    import zipfile
    zf = zipfile.ZipFile(io.BytesIO(out))
    rels_xml = zf.read("word/_rels/document.xml.rels").decode("utf-8")
    assert "https://x.example" in rels_xml


def test_filename_sanitization_drops_separators():
    assert sanitize_filename("a/b\\c.docx") == "abc.docx"


def test_filename_sanitization_caps_at_100():
    long = "x" * 200
    out = sanitize_filename(long + ".docx")
    assert len(out) <= 100


def test_filename_sanitization_falls_back_on_empty():
    assert sanitize_filename("   ") == "botnim-document.docx"
    assert sanitize_filename("") == "botnim-document.docx"


def test_multi_section_ordering_preserved():
    req = WordDocRequest(
        title="t",
        sections=[
            WordDocSection(heading="A", level=1, body_md="aaa"),
            WordDocSection(heading="B", level=2, body_md="bbb"),
            WordDocSection(heading="C", level=1, body_md="ccc"),
        ],
    )
    out = render_word_doc(req)
    doc = _docx_bytes_to_doc(out)
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading") and p.text in {"A", "B", "C"}]
    assert headings == ["A", "B", "C"]
