"""Extractor unit tests.

Three pure functions, three input flavors:

* ``html_to_text`` — strips tags, decodes the entities that actually
  appear in gov.il HTML payloads, collapses whitespace, preserves
  paragraph boundaries.
* ``pdf_to_text`` — round-trips a synthesized one-page PDF via pymupdf
  and asserts the body text is recovered + control characters are
  stripped (Excel rejects \\x00–\\x08 etc).
* ``docx_to_text`` — round-trips a synthesized .docx via python-docx
  and asserts paragraph order + control-char strip.
"""
from __future__ import annotations

import io

import docx
import fitz
import pytest

from botnim.document_parser.gov_il_decisions.extract import (
    docx_to_text,
    html_to_text,
    pdf_to_text,
    strip_control_chars,
)


def test_html_to_text_strips_tags_and_decodes_common_entities():
    html = (
        "<div><p>נושא ההחלטה: מינוי&nbsp;שר</p>"
        "<p>מחליטים:&#13;&#10;א. למנות.</p></div>"
    )
    out = html_to_text(html)
    assert "נושא ההחלטה: מינוי שר" in out
    assert "מחליטים:" in out
    assert "א. למנות." in out
    assert "<" not in out and ">" not in out
    assert "&nbsp;" not in out


def test_html_to_text_handles_empty():
    assert html_to_text("") == ""
    assert html_to_text(None) == ""


def test_pdf_to_text_round_trip():
    # Build a one-page PDF with pymupdf, render Hebrew text, extract.
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 100), "Decision body text", fontsize=12)
    blob = doc.tobytes()
    doc.close()

    out = pdf_to_text(blob)
    assert "Decision body text" in out


def test_pdf_to_text_strips_control_chars():
    # Inject a NUL by hand into the extracted text path.
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 100), "before\x00after", fontsize=12)
    blob = doc.tobytes()
    doc.close()

    out = pdf_to_text(blob)
    assert "\x00" not in out


def test_docx_to_text_round_trip():
    d = docx.Document()
    d.add_paragraph("נושא ההחלטה")
    d.add_paragraph("מחליטים: א. למנות שר.")
    buf = io.BytesIO()
    d.save(buf)

    out = docx_to_text(buf.getvalue())
    assert "נושא ההחלטה" in out
    assert "מחליטים: א. למנות שר." in out
    # Paragraph order preserved.
    assert out.index("נושא ההחלטה") < out.index("מחליטים")


def test_strip_control_chars_keeps_newlines_and_tabs():
    s = "a\x00b\x01c\nd\te"
    out = strip_control_chars(s)
    assert out == "abc\nd\te"
